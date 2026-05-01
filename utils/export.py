"""
Export utilities: Convert resume markdown text to DOCX and PDF formats.
"""

import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from fpdf import FPDF


# ─── DOCX Export ────────────────────────────────────────────────────────────

def _parse_markdown_lines(text: str) -> list[dict]:
    """
    Parse markdown text into structured line objects.
    Each line gets a type: h1, h2, h3, bullet, text, empty
    """
    lines = text.split("\n")
    parsed = []

    for line in lines:
        stripped = line.strip()

        if not stripped:
            parsed.append({"type": "empty", "content": ""})
        elif stripped.startswith("### "):
            parsed.append({"type": "h3", "content": stripped[4:]})
        elif stripped.startswith("## "):
            parsed.append({"type": "h2", "content": stripped[3:]})
        elif stripped.startswith("# "):
            parsed.append({"type": "h1", "content": stripped[2:]})
        elif stripped.startswith("- ") or stripped.startswith("* "):
            parsed.append({"type": "bullet", "content": stripped[2:]})
        elif stripped.startswith("---"):
            parsed.append({"type": "hr", "content": ""})
        else:
            parsed.append({"type": "text", "content": stripped})

    return parsed


def _add_hyperlink(paragraph, url, text):
    """
    A helper function that adds a hyperlink to a paragraph.
    """
    # This gets access to the document.xml.rels file and adds a new relation id
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color element
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(qn('w:val'), '0066CC') # Blue color
    rPr.append(c)

    # Add underline element
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Join them all together
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

import docx.oxml

def _add_formatted_text(paragraph, text: str):
    """
    Add text to a paragraph, handling **bold**, *italic*, and raw URLs.
    """
    # Pattern to find URLs
    url_pattern = r'(https?://[^\s]+)'
    
    # Split text by URLs first
    parts = re.split(url_pattern, text)
    
    for part in parts:
        if re.match(url_pattern, part):
            _add_hyperlink(paragraph, part, part)
        else:
            # Handle bold and italic within the non-URL parts
            sub_parts = re.split(r'(\*\*.*?\*\*)', part)
            for sp in sub_parts:
                if sp.startswith("**") and sp.endswith("**"):
                    run = paragraph.add_run(sp[2:-2])
                    run.bold = True
                else:
                    italic_parts = re.split(r'(\*.*?\*)', sp)
                    for ip in italic_parts:
                        if ip.startswith("*") and ip.endswith("*") and not ip.startswith("**"):
                            run = paragraph.add_run(ip[1:-1])
                            run.italic = True
                        else:
                            paragraph.add_run(ip)


def export_to_docx(resume_text: str, attachments: list = None) -> bytes:
    """
    Convert resume markdown text to a professionally formatted DOCX file.
    Appends attachments if provided.

    Args:
        resume_text: Resume content in markdown format
        attachments: List of document dicts from user_documents

    Returns:
        DOCX file as bytes
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    parsed = _parse_markdown_lines(resume_text)

    for item in parsed:
        if item["type"] == "h1":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(item["content"])
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            p.space_after = Pt(4)

        elif item["type"] == "h2":
            p = doc.add_paragraph()
            run = p.add_run(item["content"].upper())
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x6B)
            p.space_before = Pt(12)
            p.space_after = Pt(4)
            # Add a bottom border (horizontal rule effect)
            p_format = p.paragraph_format
            p_format.space_after = Pt(4)

        elif item["type"] == "h3":
            p = doc.add_paragraph()
            _add_formatted_text(p, item["content"])
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
            p.space_before = Pt(8)
            p.space_after = Pt(2)

        elif item["type"] == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, item["content"])
            for run in p.runs:
                run.font.size = Pt(10.5)
            p.space_after = Pt(1)

        elif item["type"] == "text":
            p = doc.add_paragraph()
            _add_formatted_text(p, item["content"])
            for run in p.runs:
                run.font.size = Pt(10.5)
            p.space_after = Pt(2)

        elif item["type"] == "hr":
            p = doc.add_paragraph()
            p.space_before = Pt(2)
            p.space_after = Pt(2)

        # Skip empty lines (just add spacing)

    # Handle attachments
    if attachments:
        import os
        for doc_info in attachments:
            file_path = doc_info.get("file_path")
            if not file_path or not os.path.exists(file_path):
                continue
                
            doc.add_page_break()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Attachment: {doc_info.get('title', 'Document')}")
            run.bold = True
            
            ext = file_path.lower().split('.')[-1]
            if ext in ['png', 'jpg', 'jpeg']:
                try:
                    doc.add_picture(file_path, width=Inches(6.0))
                except Exception:
                    doc.add_paragraph("[Image rendering failed]")
            elif ext == 'pdf':
                doc.add_paragraph("[PDF attached. Please refer to the PDF export for full visual inclusion or view original file.]")
            else:
                doc.add_paragraph(f"[Attached File: {doc_info.get('title')}]")

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ─── PDF Export ─────────────────────────────────────────────────────────────

class ResumePDF(FPDF):
    """Custom PDF class for resume generation."""

    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)

    def header(self):
        pass  # No header for resumes

    def footer(self):
        pass  # No footer for clean resume look


import markdown

def export_to_pdf(resume_text: str, attachments: list = None) -> bytes:
    """
    Convert resume markdown text to a clean, ATS-friendly PDF with clickable links.
    Appends attachments using pypdf if provided.

    Args:
        resume_text: Resume content in markdown format
        attachments: List of document dicts from user_documents

    Returns:
        PDF file as bytes
    """
    pdf = ResumePDF()
    pdf.add_page()
    pdf.set_margins(15, 15, 15)

    # Clean the markdown text to make it more fpdf2 write_html friendly
    # Remove excessive horizontal rules which fpdf2 struggles with
    clean_text = resume_text.replace("---", "")
    
    # Convert Markdown to HTML
    # We use a custom processor or simple regex to ensure raw URLs are clickable in PDF
    url_pattern = r'(?<!href=")(https?://[^\s<]+)'
    html_text = markdown.markdown(clean_text)
    
    # Wrap raw URLs in <a> tags for PDF
    html_content = re.sub(url_pattern, r'<a href="\1">\1</a>', html_text)

    # Basic CSS for fpdf2 to understand formatting
    html_header = """
    <style>
    h1 { text-align: center; font-family: Helvetica; font-size: 16pt; color: #1a1a2e; }
    h2 { font-family: Helvetica; font-size: 12pt; color: #2c3e6b; margin-top: 10px; }
    h3 { font-family: Helvetica; font-size: 11pt; color: #34495e; font-weight: bold; }
    p { font-family: Helvetica; font-size: 10pt; color: #333333; line-height: 1.5; }
    li { font-family: Helvetica; font-size: 10pt; color: #333333; }
    a { color: #0066cc; text-decoration: none; }
    </style>
    """

    
    full_html = html_header + html_content
    
    try:
        pdf.write_html(full_html)
    except Exception as e:
        # Fallback to simple text if HTML parsing fails
        pdf.set_font("Helvetica", size=10)
        pdf.multi_cell(0, 5, resume_text.encode('latin-1', 'replace').decode('latin-1'))
        
    # Save to bytes
    buffer = io.BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    
    # Handle attachments
    if attachments:
        try:
            from pypdf import PdfWriter, PdfReader
            import os
            
            writer = PdfWriter()
            # Add the generated resume PDF
            writer.append(buffer)
            
            # Add attachments
            for doc_info in attachments:
                file_path = doc_info.get("file_path")
                if not file_path or not os.path.exists(file_path):
                    continue
                    
                ext = file_path.lower().split('.')[-1]
                if ext == 'pdf':
                    try:
                        reader = PdfReader(file_path)
                        writer.append(reader)
                    except Exception:
                        pass
                elif ext in ['png', 'jpg', 'jpeg']:
                    # Create a temporary PDF for the image
                    img_pdf = FPDF()
                    img_pdf.add_page()
                    # Add title
                    img_pdf.set_font("Helvetica", 'B', 12)
                    img_pdf.cell(0, 10, f"Attachment: {doc_info.get('title', 'Document')}", ln=True, align='C')
                    try:
                        img_pdf.image(file_path, x=15, y=30, w=180)
                    except Exception:
                        img_pdf.cell(0, 10, "[Image Error]", ln=True, align='C')
                    
                    img_buf = io.BytesIO()
                    img_pdf.output(img_buf)
                    img_buf.seek(0)
                    writer.append(PdfReader(img_buf))
                    
            final_buffer = io.BytesIO()
            writer.write(final_buffer)
            final_buffer.seek(0)
            return final_buffer.getvalue()
        except ImportError:
            # If pypdf is missing, just return the main PDF
            pass
            
    return buffer.getvalue()
