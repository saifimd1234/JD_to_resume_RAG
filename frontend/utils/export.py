"""
Export utilities: Convert resume markdown text to DOCX and PDF formats.
"""

import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF


# ─── DOCX Export ────────────────────────────────────────────────────────────

def _parse_markdown_lines(text: str) -> list[dict]:
    """
    Parse markdown text into structured line objects.
    Each line gets a type: h1, h2, h3, bullet, text, empty
    """
    lines = text.split("\n")
    parsed = []

    for line in lines:
        stripped = line.strip()

        if not stripped:
            parsed.append({"type": "empty", "content": ""})
        elif stripped.startswith("### "):
            parsed.append({"type": "h3", "content": stripped[4:]})
        elif stripped.startswith("## "):
            parsed.append({"type": "h2", "content": stripped[3:]})
        elif stripped.startswith("# "):
            parsed.append({"type": "h1", "content": stripped[2:]})
        elif stripped.startswith("- ") or stripped.startswith("* "):
            parsed.append({"type": "bullet", "content": stripped[2:]})
        elif stripped.startswith("---"):
            parsed.append({"type": "hr", "content": ""})
        else:
            parsed.append({"type": "text", "content": stripped})

    return parsed


def _add_formatted_text(paragraph, text: str):
    """
    Add text to a paragraph, handling **bold** and *italic* markdown.
    """
    # Split on bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle italic within non-bold parts
            italic_parts = re.split(r'(\*.*?\*)', part)
            for ip in italic_parts:
                if ip.startswith("*") and ip.endswith("*") and not ip.startswith("**"):
                    run = paragraph.add_run(ip[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(ip)


def export_to_docx(resume_text: str) -> bytes:
    """
    Convert resume markdown text to a professionally formatted DOCX file.

    Args:
        resume_text: Resume content in markdown format

    Returns:
        DOCX file as bytes
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    parsed = _parse_markdown_lines(resume_text)

    for item in parsed:
        if item["type"] == "h1":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(item["content"])
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            p.space_after = Pt(4)

        elif item["type"] == "h2":
            p = doc.add_paragraph()
            run = p.add_run(item["content"].upper())
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x6B)
            p.space_before = Pt(12)
            p.space_after = Pt(4)
            # Add a bottom border (horizontal rule effect)
            p_format = p.paragraph_format
            p_format.space_after = Pt(4)

        elif item["type"] == "h3":
            p = doc.add_paragraph()
            _add_formatted_text(p, item["content"])
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
            p.space_before = Pt(8)
            p.space_after = Pt(2)

        elif item["type"] == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, item["content"])
            for run in p.runs:
                run.font.size = Pt(10.5)
            p.space_after = Pt(1)

        elif item["type"] == "text":
            p = doc.add_paragraph()
            _add_formatted_text(p, item["content"])
            for run in p.runs:
                run.font.size = Pt(10.5)
            p.space_after = Pt(2)

        elif item["type"] == "hr":
            p = doc.add_paragraph()
            p.space_before = Pt(2)
            p.space_after = Pt(2)

        # Skip empty lines (just add spacing)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ─── PDF Export ─────────────────────────────────────────────────────────────

class ResumePDF(FPDF):
    """Custom PDF class for resume generation."""

    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)

    def header(self):
        pass  # No header for resumes

    def footer(self):
        pass  # No footer for clean resume look


def export_to_pdf(resume_text: str) -> bytes:
    """
    Convert resume markdown text to a clean, ATS-friendly PDF.

    Args:
        resume_text: Resume content in markdown format

    Returns:
        PDF file as bytes
    """
    pdf = ResumePDF()
    pdf.add_page()
    pdf.set_margins(18, 15, 18)

    # Use built-in fonts (ATS-friendly, no encoding issues)
    parsed = _parse_markdown_lines(resume_text)

    for item in parsed:
        if item["type"] == "h1":
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_text_color(26, 26, 46)
            pdf.cell(0, 10, item["content"], ln=True, align="C")
            pdf.ln(2)

        elif item["type"] == "h2":
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(44, 62, 107)
            pdf.cell(0, 8, item["content"].upper(), ln=True)
            # Draw line under section header
            pdf.set_draw_color(44, 62, 107)
            pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 174, pdf.get_y())
            pdf.ln(3)

        elif item["type"] == "h3":
            pdf.set_font("Helvetica", "B", 10.5)
            pdf.set_text_color(52, 73, 94)
            # Strip markdown bold markers for PDF
            clean_text = item["content"].replace("**", "")
            pdf.cell(0, 7, clean_text, ln=True)
            pdf.ln(1)

        elif item["type"] == "bullet":
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(51, 51, 51)
            # Clean markdown formatting
            clean_text = item["content"].replace("**", "")
            pdf.cell(6, 5, "-")
            pdf.multi_cell(0, 5, clean_text)
            pdf.ln(1)

        elif item["type"] == "text":
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(51, 51, 51)
            clean_text = item["content"].replace("**", "")
            pdf.multi_cell(0, 5, clean_text)
            pdf.ln(1)

        elif item["type"] == "hr":
            pdf.ln(2)

        elif item["type"] == "empty":
            pdf.ln(2)

    # Save to bytes
    buffer = io.BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer.getvalue()
